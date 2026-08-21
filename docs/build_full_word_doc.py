import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Color Palette Constants
    COLOR_PRIMARY = RGBColor(0x02, 0x84, 0xC7)      # Cyan / Blue #0284c7
    COLOR_SECONDARY = RGBColor(0x4F, 0x46, 0xE5)    # Indigo #4f46e5
    COLOR_DARK = RGBColor(0x0F, 0x17, 0x2A)         # Slate Dark #0f172a
    COLOR_TEXT = RGBColor(0x33, 0x41, 0x55)         # Body Text #334155
    COLOR_MUTED = RGBColor(0x64, 0x74, 0x8B)        # Muted Text #64748b
    COLOR_EMERALD = RGBColor(0x05, 0x96, 0x69)      # Emerald Green #059669

    HEX_PRIMARY = "0284C7"
    HEX_SECONDARY = "4F46E5"
    HEX_LIGHT_BG = "F8FAFC"
    HEX_CARD_BG = "F0F9FF"
    HEX_EMERALD_BG = "ECFDF5"
    HEX_BORDER = "E2E8F0"
    HEX_HEADER_BG = "0F172A"

    # Style Helpers
    def style_heading_1(p, text):
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(18)
        run.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def style_heading_2(p, text):
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return p

    def style_heading_3(p, text):
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = COLOR_DARK
        return p

    def add_body_p(doc, text, bold_prefix=None, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.font.name = "Calibri"
            run_b.font.size = Pt(10.5)
            run_b.bold = True
            run_b.font.color.rgb = COLOR_DARK
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.italic = italic
        run.font.color.rgb = COLOR_TEXT
        return p

    def add_bullet(doc, bold_title, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_title:
            run_b = p.add_run(bold_title)
            run_b.font.name = "Calibri"
            run_b.font.size = Pt(10)
            run_b.bold = True
            run_b.font.color.rgb = COLOR_DARK
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_TEXT
        return p

    def set_cell_background(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'''
            <w:tcMar {nsdecls("w")}>
                <w:top w:w="{top}" w:type="dxa"/>
                <w:bottom w:w="{bottom}" w:type="dxa"/>
                <w:left w:w="{left}" w:type="dxa"/>
                <w:right w:w="{right}" w:type="dxa"/>
            </w:tcMar>
        ''')
        tcPr.append(tcMar)

    def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
        tblPr = table._tbl.tblPr
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
                <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideV w:val="none"/>
            </w:tblBorders>
        ''')
        tblPr.append(borders)

    def add_callout(doc, text, title="NOTE", border_color="0284C7", bg_color="F0F9FF"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(f'''
            <w:tcBorders {nsdecls("w")}>
                <w:top w:val="none"/>
                <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>
                <w:bottom w:val="none"/>
                <w:right w:val="none"/>
            </w:tcBorders>
        ''')
        tcPr.append(borders)

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run_title = p.add_run(f"📌 {title}: ")
        run_title.bold = True
        run_title.font.name = "Calibri"
        run_title.font.size = Pt(10)
        run_title.font.color.rgb = RGBColor.from_string(border_color)

        run_text = p.add_run(text)
        run_text.font.name = "Calibri"
        run_text.font.size = Pt(9.5)
        run_text.font.color.rgb = COLOR_TEXT

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_code_block(doc, code_str):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=120, bottom=120, left=160, right=160)

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(code_str)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x7D, 0xD3, 0xFC)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def create_styled_table(doc, headers, data, col_widths=None):
        table = doc.add_table(rows=len(data) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_borders(table, color="CBD5E1", sz="4", val="single")

        # Header Row
        hdr_row = table.rows[0]
        tblPr = table._tbl.tblPr
        for i, header in enumerate(headers):
            cell = hdr_row.cells[i]
            if col_widths and i < len(col_widths):
                cell.width = Inches(col_widths[i])
            set_cell_background(cell, "1E293B")
            set_cell_margins(cell, top=140, bottom=140, left=140, right=140)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(header)
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Data Rows
        for r_idx, row_data in enumerate(data):
            row = table.rows[r_idx + 1]
            bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, cell_value in enumerate(row_data):
                cell = row.cells[c_idx]
                if col_widths and c_idx < len(col_widths):
                    cell.width = Inches(col_widths[c_idx])
                set_cell_background(cell, bg_color)
                set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(str(cell_value))
                run.font.name = "Calibri"
                run.font.size = Pt(9)
                run.font.color.rgb = COLOR_TEXT

        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        return table

    # =========================================================================
    # COVER / TITLE PAGE
    # =========================================================================
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(72)

    p_badge = doc.add_paragraph()
    p_badge.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_badge = p_badge.add_run("ENTERPRISE PLATFORM SPECIFICATION & OPERATING GUIDE")
    run_badge.font.name = "Calibri"
    run_badge.font.size = Pt(11)
    run_badge.bold = True
    run_badge.font.color.rgb = COLOR_PRIMARY

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(8)
    run_title = p_title.add_run("OntoForge (Quick-Pasteur)")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(32)
    run_title.bold = True
    run_title.font.color.rgb = COLOR_DARK

    p_subtitle = doc.add_paragraph()
    p_subtitle.paragraph_format.space_before = Pt(0)
    p_subtitle.paragraph_format.space_after = Pt(20)
    run_sub = p_subtitle.add_run("Enterprise Relational-to-Graph & W3C OWL 2.0 Ontology Automated Transformation Platform")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(15)
    run_sub.font.color.rgb = COLOR_SECONDARY

    p_rule = doc.add_paragraph()
    p_rule.paragraph_format.space_after = Pt(24)
    run_rule = p_rule.add_run("―" * 48)
    run_rule.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.line_spacing = 1.25
    p_meta.paragraph_format.space_after = Pt(40)
    
    runs_meta = [
        ("Platform Version: ", True), ("1.0.0 Enterprise Release\n", False),
        ("Standard Compliance: ", True), ("W3C OWL 2.0 DL, RDF 1.1, RDFS, SKOS, SPARQL 1.1, OpenCypher\n", False),
        ("Document Classification: ", True), ("Technical Architecture, Functional Specification & Operational Manual\n", False),
        ("Document Status: ", True), ("Approved / Complete\n", False),
        ("Target Audience: ", True), ("Enterprise Architects, Knowledge Engineers, Data Scientists, Ontologists, DBAs", False),
    ]
    for m_text, m_bold in runs_meta:
        r = p_meta.add_run(m_text)
        r.font.name = "Calibri"
        r.font.size = Pt(10)
        r.bold = m_bold
        r.font.color.rgb = COLOR_DARK if m_bold else COLOR_TEXT

    doc.add_page_break()

    # =========================================================================
    # SECTION 1: EXECUTIVE SUMMARY & STRATEGIC OVERVIEW
    # =========================================================================
    p1 = doc.add_paragraph()
    style_heading_1(p1, "1. Executive Summary & Strategic Overview")

    add_body_p(doc, "In modern enterprise ecosystems, business data remains fragmented across heterogeneous relational database management systems (RDBMS) such as PostgreSQL, Microsoft SQL Server, Oracle, MySQL, and cloud data warehouses like Snowflake and Databricks. While relational architectures excel at structured transaction processing, they inherently struggle with contextual knowledge representation, semantic reasoning, cross-domain relationship discovery, and automated inference.")

    add_body_p(doc, "OntoForge (Quick-Pasteur Enterprise Engine) bridges this enterprise gap by providing an end-to-end automated platform that ingests relational metadata catalogs, introspects schema topology, analyzes data distributions and privacy characteristics, embeds corporate business governance rules, and synthesizes standards-compliant W3C OWL 2.0 Web Ontology Language models and labeled Property Knowledge Graphs.")

    add_callout(doc, "OntoForge transforms physical tables, primary keys, and foreign keys into rich ontological classes (owl:Class), scalar datatype properties (owl:DatatypeProperty), and bidirectional relationship axioms (owl:ObjectProperty, owl:inverseOf) with functional primary key articulation (owl:hasKey), ready for instantaneous deployment to knowledge graphs or semantic graph reasoning triplestores.", title="STRATEGIC VALUE PROPOSITION", border_color="0284C7", bg_color="F0F9FF")

    p1_sub = doc.add_paragraph()
    style_heading_2(p1_sub, "1.1 Key Enterprise Value Drivers")

    add_bullet(doc, "Automated Semantic Lifting: ", "Eliminates months of manual ontology engineering by automatically converting database tables, columns, and foreign key constraints into semantically typed OWL 2.0 DL models.")
    add_bullet(doc, "Multi-Database Source Consolidation: ", "Simultaneously ingests from multiple disconnected databases into an isolated, project-scoped enterprise knowledge graph.")
    add_bullet(doc, "Active Business Rules Harmonization: ", "Allows business analysts and data stewards to express domain governance rules in plain English, which are compiled into ontology axioms and RDFS annotations.")
    add_bullet(doc, "High-Precision Taxonomy Classification: ", "Categorizes relational objects into standard upper-ontology classifications: MasterEntity, TransactionalEntity, ReferenceEntity, and AssociativeEntity.")
    add_bullet(doc, "Multi-Format Export Ecosystem: ", "One-click export to W3C Turtle (.ttl), OWL/XML (.owl), Neo4j Cypher DDL scripts (.cypher), and GraphML (.graphml).")
    add_bullet(doc, "Stateless In-Memory Ontology Sandbox: ", "Provides a zero-persistence exploration environment allowing users to upload, validate, parse, and visually inspect external RDF/OWL files in real time.")

    # =========================================================================
    # SECTION 2: SYSTEM ARCHITECTURE & TECHNOLOGY STACK
    # =========================================================================
    doc.add_page_break()
    p2 = doc.add_paragraph()
    style_heading_1(p2, "2. System Architecture & Technical Specifications")

    add_body_p(doc, "OntoForge is engineered using a decoupled, high-performance architecture consisting of a Python FastAPI backend service and an executive light-mode Single Page Application (SPA) frontend powered by Cytoscape.js and native ES6+ modules.")

    p2_arch = doc.add_paragraph()
    style_heading_2(p2_arch, "2.1 High-Level End-to-End Architectural Model")

    # Embed High-Resolution Graphic Diagram
    img_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "architecture_diagram.png")
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(6.5))
        
        p_caption = doc.add_paragraph()
        p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_caption.paragraph_format.space_after = Pt(12)
        run_cap = p_caption.add_run("Figure 2.1: OntoForge End-to-End Enterprise Architecture, Ingestion & Semantic Synthesis Flow")
        run_cap.font.name = "Calibri"
        run_cap.font.size = Pt(9)
        run_cap.italic = True
        run_cap.font.color.rgb = COLOR_MUTED

    p2_layers = doc.add_paragraph()
    style_heading_3(p2_layers, "2.1.1 Architectural Layer Decomposition & Interaction Flow")

    arch_flow_data = [
        ["Layer / Subsystem", "Primary Capabilities", "Inputs & Predecessors", "Outputs & Handoffs"],
        ["Layer 1: Multi-Source Ingestion", "PostgreSQL, MSSQL, MySQL, Oracle, SQLite, Snowflake & Databricks connection pool & health diagnostics", "Raw Database Endpoints & Encrypted Credentials", "Live Connection Handlers & Schemas"],
        ["Layer 2: Discovery & Profiling", "Information Schema discovery, SQL-to-XSD datatype mapping, automated PII detection & quality scoring", "Relational Database Catalogs", "Physical Metadata Catalog & Column Profiles"],
        ["Layer 3: Business Governance", "Plain-English governance rules engine (Validation, Masking, Transformation, Lookup, Quality)", "Business Analyst & Data Steward Inputs", "Rule Bindings & Target Annotations"],
        ["Layer 4: Semantic OWL Synthesis", "W3C OWL 2.0 DL generator, upper-taxonomies, Datatype/Object properties, owl:hasKey, owl:inverseOf", "Metadata Catalog + Bound Business Rules", "RDF Triples & Formal OWL 2.0 DL Ontology"],
        ["Layer 5: Interactive Visualizer", "Cytoscape.js canvas, Timbr design system, multi-mode view switcher, root path filter, stateless sandbox", "Compiled OWL Model & External RDF Files", "Interactive Concept & Mapping Diagrams"],
        ["Layer 6: Target Graph & Export", "W3C Turtle (.ttl), OWL/XML (.owl), Neo4j Cypher DDL (.cypher), GraphML (.graphml), direct Bolt sync", "Synthesized Ontology & Property Graph", "Target Knowledge Graph (Neo4j / Memgraph)"]
    ]
    create_styled_table(doc, arch_flow_data[0], arch_flow_data[1:], col_widths=[1.5, 2.2, 1.4, 1.4])

    p2_tech = doc.add_paragraph()
    style_heading_2(p2_tech, "2.2 Technology Stack Matrix")

    tech_data = [
        ["Subsystem", "Component / Framework", "Version", "Role & Responsibility"],
        ["Backend Core", "FastAPI (ASGI)", "0.115+", "High-speed asynchronous REST API, OpenAPI docs & routing"],
        ["ASGI Server", "Uvicorn", "0.32+", "Production ASGI server with multiprocessing & live-reload"],
        ["ORM & DB Engine", "SQLAlchemy", "2.0+", "Relational database mapping, connection pooling & schema sync"],
        ["Application DB", "SQLite / PostgreSQL", "3.x / 15+", "Stores metadata catalogs, project isolation states & rules"],
        ["Semantic Web", "RDFLib", "7.0+", "W3C RDF graph synthesis, OWL 2.0 DL axiomatization & serialization"],
        ["Graph Engine", "NetworkX", "3.4+", "Topological analysis, cycle detection & GraphML generation"],
        ["Validation", "Pydantic v2", "2.10+", "Strict runtime request/response serialization & schema integrity"],
        ["Frontend Shell", "Vanilla HTML5 / ES6+", "ES2022+", "Single Page Application (SPA) modular UI controllers"],
        ["Graph Canvas", "Cytoscape.js", "3.28+", "Force-directed physics graph simulation (CoSE, Breadth-First)"],
        ["Typography", "Google Inter / JetBrains", "WebFonts", "Executive enterprise dashboard visual hierarchy"],
        ["Target Sync", "Neo4j Bolt Driver / REST", "5.x+", "Direct live Cypher constraint & node/edge synchronization"]
    ]
    create_styled_table(doc, tech_data[0], tech_data[1:], col_widths=[1.2, 1.6, 0.9, 2.8])

    # =========================================================================
    # SECTION 3: CORE FUNCTIONAL MODULES
    # =========================================================================
    doc.add_page_break()
    p3 = doc.add_paragraph()
    style_heading_1(p3, "3. Core Functional Modules & Capabilities")

    # Module 1
    p3_1 = doc.add_paragraph()
    style_heading_2(p3_1, "3.1 Multi-Source Database Connectors & Target Graph Topology")
    add_body_p(doc, "The Connectors subsystem enables projects to map one or more heterogeneous relational databases into a consolidated workspace. Connection definitions are completely isolated per project, ensuring multi-tenant data safety.")
    
    add_bullet(doc, "Supported Connectors: ", "PostgreSQL, Microsoft SQL Server, MySQL, MariaDB, Oracle Database, SQLite, Snowflake Data Cloud, and Databricks Lakehouse.")
    add_bullet(doc, "Connection Verification: ", "Integrated connection testing verifies host reachability, authentication credentials, and database permissions prior to schema discovery.")
    add_bullet(doc, "Credential Security: ", "Database connection strings and passwords are encrypted and never exposed in raw client logs.")
    add_bullet(doc, "Unified Target Destination: ", "Supports configuration of a single target knowledge graph instance (Neo4j, Memgraph, Apache AGE, AWS Neptune) with Bolt protocol synchronization.")

    # Module 2
    p3_2 = doc.add_paragraph()
    style_heading_2(p3_2, "3.2 Automated Metadata Discovery & Schema Profiling")
    add_body_p(doc, "The Discovery engine inspects the physical relational catalog of connected databases to construct a structured metadata repository without requiring manual schema definitions.")

    add_bullet(doc, "Catalog Introspection: ", "Extracts table names, views, column definitions, nullability flags, comments, primary keys, and explicit foreign key relationships.")
    add_bullet(doc, "SQL-to-XSD Type Translation: ", "Converts SQL dialect types (e.g. VARCHAR, INT, BIGINT, NUMERIC, TIMESTAMP, BYTEA) to standard W3C XSD datatypes (xsd:string, xsd:integer, xsd:decimal, xsd:dateTime, xsd:hexBinary).")
    add_bullet(doc, "Domain Type Classification: ", "Heuristically classifies discovered tables into Dimension/Master, Fact/Transactional, Reference/Lookup, or Associative/Bridge domains based on naming conventions and foreign key cardinality.")
    add_bullet(doc, "PII Privacy Detection: ", "Scans column headers and sample distributions to identify Personally Identifiable Information (EMAIL, SSN, PHONE, CREDIT_CARD, NAME).")
    add_bullet(doc, "Statistical Data Profiling: ", "Calculates table row counts, distinct value counts, null ratios, and data completeness scores.")

    # Module 3
    p3_3 = doc.add_paragraph()
    style_heading_2(p3_3, "3.3 Business & Governance Rules Engine")
    add_body_p(doc, "OntoForge empowers non-technical domain experts to define plain-English corporate governance policies that are automatically codified into the synthesized W3C ontology.")

    rule_categories = [
        ["Rule Type", "Purpose & Semantics", "Ontological Realization"],
        ["VALIDATION", "Enforces value bounds, range constraints, and regex formatting", "Annotated as eonto:hasBusinessRule and owl:Restriction axioms"],
        ["TRANSFORMATION", "Specifies field derivation, unit conversion, or currency normalization", "Articulated in RDFS comments and mapping execution specs"],
        ["LOOKUP", "Enforces reference dictionary mappings and valid enum codes", "Binds class attributes to eonto:ReferenceEntity taxonomy instances"],
        ["MASKING", "Defines data anonymization and privacy obfuscation policies", "Tags target properties with PII security governance annotations"],
        ["QUALITY", "Specifies threshold criteria for data freshness, nullability, and completeness", "Embedded into table profiling records and class metadata"],
        ["ENRICHMENT", "Defines synthetic semantic attributes or external ontology links", "Synthesizes supplementary Datatype and Object properties"],
        ["CUSTOM", "Captures bespoke enterprise business logic and compliance standards", "Custom RDF annotations bound to target owl:Class definitions"]
    ]
    create_styled_table(doc, rule_categories[0], rule_categories[1:], col_widths=[1.5, 2.5, 2.5])

    # Module 4
    doc.add_page_break()
    p3_4 = doc.add_paragraph()
    style_heading_2(p3_4, "3.4 Semantic Web & W3C OWL 2.0 DL Ontology Engine")
    add_body_p(doc, "The OWL Generation engine (powered by RDFLib 7.0+) performs rigorous translation of relational metadata into a W3C OWL 2.0 Description Logic (DL) compliant ontology. It implements industry-standard semantic web practices to ensure full interoperability with Protégé, Apache Jena, GraphDB, and TopBraid Composer.")

    add_bullet(doc, "Singular PascalCase Concepts: ", "Converts plural snake_case table names (e.g. customer_orders, proteins, biological_assays) into standard singular PascalCase OWL classes (CustomerOrder, Protein, BiologicalAssay).")
    add_bullet(doc, "Upper-Level Taxonomy Hierarchy: ", "Rooted in standard taxonomies subClassOf owl:Thing (eonto:MasterEntity, eonto:TransactionalEntity, eonto:ReferenceEntity, eonto:AssociativeEntity).")
    add_bullet(doc, "Functional Primary Key Articulation: ", "Declares primary key attributes as owl:FunctionalProperty, annotates with eonto:isPrimaryKey, and binds them to the enclosing class using W3C OWL 2 owl:hasKey RDF List collections.")
    add_bullet(doc, "Domain-Specific Object Properties: ", "Infers meaningful relationship names based on foreign key semantics (e.g. placedBy / hasOrder, bindsCompound / boundByProtein, billsProduct / billedInInvoiceItem) rather than generic relatesTo prefixes.")
    add_bullet(doc, "Bidirectional Inverse Properties: ", "Automatically synthesizes forward and backward relationship pairs interconnected via formal owl:inverseOf axioms.")

    add_code_block(doc, 
        "# Sample Generated W3C Turtle (.ttl) Excerpt\n"
        "@prefix eonto: <http://enterprise.org/ontology#> .\n"
        "@prefix owl: <http://www.w3.org/2002/07/owl#> .\n"
        "@prefix rdfs: <http://www.w3.org/2000/01/rdf-schema#> .\n"
        "@prefix xsd: <http://www.w3.org/2001/XMLSchema#> .\n\n"
        "eonto:Protein a owl:Class ;\n"
        "    rdfs:label \"Protein\" ;\n"
        "    rdfs:subClassOf eonto:MasterEntity ;\n"
        "    eonto:hasPrimaryKey \"proteinId\" ;\n"
        "    owl:hasKey ( eonto:proteinId ) .\n\n"
        "eonto:bindsCompound a owl:ObjectProperty, owl:FunctionalProperty ;\n"
        "    rdfs:domain eonto:Protein ;\n"
        "    rdfs:range eonto:ChemicalCompound ;\n"
        "    owl:inverseOf eonto:boundByProtein .\n"
    )

    # Module 5
    p3_5 = doc.add_paragraph()
    style_heading_2(p3_5, "3.5 Interactive Graphical Ontology Visualizer (Timbr.ai Design System)")
    add_body_p(doc, "The Graphical Ontology interface provides an interactive, visual representation of the synthesized ontology. Drawing inspiration from modern semantic modeling environments (such as Timbr.ai), it renders classes, attributes, and relationships with clear visual differentiation.")

    add_bullet(doc, "Visual Encoding: ", "Classes are rendered as sky-blue circular concept nodes, superclass taxonomies as navy circles, datatype properties as teal squares, and object properties as emerald relationship diamonds.")
    add_bullet(doc, "Triple Mode Switcher: ", "Enables instant switching between 'Semantic Ontology' (concept model), 'Source Metadata' (physical table view), and 'Mapping View' (end-to-end visual lineage flows connecting source columns to ontology attributes).")
    add_bullet(doc, "Multi-Class Taxonomy Filter: ", "Allows selecting one or more classes to trace and isolate their complete inheritance lineage paths up to owl:Thing.")
    add_bullet(doc, "Interactive Right-Side Details Drawer: ", "Clicking any node opens a comprehensive inspection drawer displaying physical table lineage, scalar attributes, active business rules, and interconnected relationships with one-click subclass creation.")

    # Module 6
    p3_6 = doc.add_paragraph()
    style_heading_2(p3_6, "3.6 Stateless Upload & Sandbox Viewer")
    add_body_p(doc, "The Stateless Sandbox allows enterprise users to drag and drop external RDF/OWL files (.ttl, .owl, .rdf, .xml, .jsonld, .nt) or paste raw Turtle markup to immediately parse, validate, and visually explore the model in-memory with zero database writes.")

    add_bullet(doc, "Multi-Format Ingestion: ", "Auto-detects and parses Turtle, RDF/XML, JSON-LD, and N-Triples formats using in-memory RDFLib streams.")
    add_bullet(doc, "Interactive Sub-Tabs: ", "Features 4 dedicated sub-views: (1) Knowledge Graph with physics-based layout engine selector (CoSE, Concentric, Breadth-First, Circle, Grid), (2) OWL Classes Grid, (3) Properties & Relationships Table, and (4) Formatted Raw Turtle Source Viewer with copy and download utilities.")
    add_bullet(doc, "Preset Demonstrations: ", "Includes pre-loaded biological/assay knowledge graphs (Pasteur Model) and e-commerce supply chain graphs for rapid evaluation.")

    # =========================================================================
    # SECTION 4: COMPLETE USER OPERATING MANUAL
    # =========================================================================
    doc.add_page_break()
    p4 = doc.add_paragraph()
    style_heading_1(p4, "4. Step-by-Step User Operating Manual")

    add_body_p(doc, "This section outlines the standard end-to-end workflow for discovering, modeling, governing, and exporting enterprise ontologies and knowledge graphs.")

    steps = [
        ("Step 1: Launching the Application & Project Management", 
         "1. Navigate to http://127.0.0.1:8000/app in any modern web browser.\n"
         "2. To create a new workspace, click '➕ New Project' in the top navigation bar.\n"
         "3. Enter a Project Name (e.g. 'Enterprise Healthcare Ontology'), unique Project Code (e.g. 'HEALTH_2026'), and optional description.\n"
         "4. Select the active project from the dropdown. Switching projects immediately resets and refreshes all downstream tabs."),

        ("Step 2: Adding Source Database Connectors",
         "1. Open the '🔌 Database Connectors' tab on the left sidebar.\n"
         "2. Click '➕ Add Source Database'.\n"
         "3. Select the Connector Type (e.g. PostgreSQL, SQL Server, MySQL, SQLite, Snowflake).\n"
         "4. Enter the Host, Port, Database Name, Username, and Password.\n"
         "5. Click 'Test Connection' to verify connectivity. Once verified, click 'Save Connector'."),

        ("Step 3: Configuring the Target Knowledge Graph",
         "1. On the Connectors tab, click '🎯 Configure Target Graph DB' on the right panel.\n"
         "2. Choose the Target Graph Engine (Neo4j, Memgraph, Apache AGE, or AWS Neptune).\n"
         "3. Specify the Bolt/HTTP connection endpoint (e.g. bolt://localhost:7687) and credentials.\n"
         "4. Save configuration to establish the target graph sync endpoint."),

        ("Step 4: Executing Automated Metadata Discovery",
         "1. Navigate to the '🔍 Metadata Discovery' tab.\n"
         "2. Select a specific source connector or choose 'All Mapped Source Connectors'.\n"
         "3. Click '⚡ Run Auto Discovery'. The platform introspects the connected databases, populating the table with discovered schemas, primary keys, and inferred domains.\n"
         "4. Review and adjust table classifications (Master, Transactional, Reference, Associative) as needed."),

        ("Step 5: Inspecting Data Quality & Privacy Profiling",
         "1. Navigate to '📈 Data Profiling & Quality'.\n"
         "2. Click '⚡ Run Data Profiling' to calculate column-level statistics, distinct value distributions, and null ratios.\n"
         "3. Review the Automated PII Tag badges (e.g. EMAIL, SSN, PHONE) to ensure compliance with enterprise privacy standards."),

        ("Step 6: Defining Business & Governance Rules",
         "1. Navigate to '⚙️ Business Rules Engine'.\n"
         "2. Click '➕ Add Business Rule'.\n"
         "3. Select the Rule Type (Validation, Transformation, Masking, Quality, Lookup, Enrichment).\n"
         "4. Target the specific entity (e.g. 'Patient' or 'Order') and attribute (e.g. 'totalAmount' or 'email').\n"
         "5. Write the plain-English rule specification (e.g. 'Order total amount must be strictly greater than zero'). Click 'Save Rule'."),

        ("Step 7: Generating & Customizing the W3C OWL Ontology",
         "1. Navigate to '🧠 OWL Ontology Editor'.\n"
         "2. Click '🔄 Re-Generate Ontology'. The system compiles the metadata and business rules into W3C OWL 2.0 DL.\n"
         "3. Inspect each class card to review its superclass taxonomy, primary key, datatype attributes, and object relationships.\n"
         "4. Use the inline editor to modify class comments, rename attributes, or add custom properties."),

        ("Step 8: Exploring in Graphical Ontology & Mapping Modes",
         "1. Navigate to '🌐 Graphical Ontology'.\n"
         "2. Use the mode toggle at the top to switch between 'Semantic Ontology', 'Source Metadata', and 'Mapping View'.\n"
         "3. Use the search bar to locate specific entities or use chip filters to isolate Classes, Properties, or Relationships.\n"
         "4. Click any node to open the Right-Side Inspection Drawer for deep-dive lineage analysis and instant subclass creation."),

        ("Step 9: Utilizing the Stateless Upload & Sandbox Viewer",
         "1. Navigate to '📤 Upload & View Ontology'.\n"
         "2. Drag and drop any external .ttl, .owl, .rdf, .xml, .jsonld, or .nt file into the dropzone (or paste raw code).\n"
         "3. Click '⚡ Parse & Visualize Ontology'. Explore the 4 sub-views (Knowledge Graph, Classes Grid, Properties Table, Raw Turtle Source).\n"
         "4. Download sanitized W3C Turtle exports or copy source markup directly to the clipboard."),

        ("Step 10: Exporting Artifacts & Syncing with Target Graph",
         "1. From the Ontology Editor or Graphical Ontology view, click '📥 Turtle (.ttl)' or '📥 OWL/XML (.owl)' to download standardized ontology files.\n"
         "2. Navigate to '🕸️ Knowledge Graph' and click '⚡ Export Cypher (.cypher)' to generate Neo4j constraint and node creation scripts.\n"
         "3. Click '🚀 Export & Sync to Target DB' to execute direct Bolt synchronization into the live target graph database.")
    ]

    for title, desc in steps:
        p_step = doc.add_paragraph()
        style_heading_2(p_step, title)
        add_body_p(doc, desc)

    # =========================================================================
    # SECTION 5: REST API SPECIFICATION
    # =========================================================================
    doc.add_page_break()
    p5 = doc.add_paragraph()
    style_heading_1(p5, "5. Comprehensive REST API Reference")

    add_body_p(doc, "OntoForge provides a fully documented RESTful API conforming to OpenAPI 3.1 standards. Interactive Swagger documentation is accessible at /docs, and ReDoc documentation is available at /redoc.")

    api_endpoints = [
        ["HTTP Method", "Endpoint Path", "Tags / Subsystem", "Description & Return Model"],
        ["POST", "/api/v1/auth/token", "Authentication", "Authenticates user and returns JWT bearer token"],
        ["GET", "/api/v1/auth/me", "Authentication", "Retrieves current authenticated user profile"],
        ["POST", "/api/v1/projects", "Projects", "Creates new isolated enterprise project"],
        ["GET", "/api/v1/projects", "Projects", "Lists all active enterprise projects"],
        ["PUT", "/api/v1/projects/{id}", "Projects", "Updates project metadata and status"],
        ["DELETE", "/api/v1/projects/{id}", "Projects", "Cascading deletion of project and associated models"],
        ["POST", "/api/v1/projects/{id}/source-connections", "Connectors", "Creates new RDBMS source connection"],
        ["GET", "/api/v1/projects/{id}/source-connections", "Connectors", "Lists all mapped source connections"],
        ["POST", "/api/v1/projects/{id}/source-connections/{cid}/test", "Connectors", "Executes live connection health check"],
        ["POST", "/api/v1/projects/{id}/graph-configs", "Connectors", "Configures target knowledge graph connection"],
        ["POST", "/api/v1/projects/{id}/metadata/discover", "Discovery", "Triggers automated schema metadata discovery"],
        ["GET", "/api/v1/projects/{id}/metadata/tables", "Discovery", "Retrieves discovered tables, columns & foreign keys"],
        ["POST", "/api/v1/projects/{id}/profiling/run", "Profiling", "Executes data profiling, null analysis & PII tagging"],
        ["GET", "/api/v1/projects/{id}/profiling/results", "Profiling", "Retrieves profiling statistics and quality scores"],
        ["POST", "/api/v1/projects/{id}/rules", "Business Rules", "Creates new business governance rule"],
        ["GET", "/api/v1/projects/{id}/rules", "Business Rules", "Lists all active and inactive business rules"],
        ["GET", "/api/v1/projects/{id}/ontology/generate", "Ontology Engine", "Synthesizes W3C OWL 2.0 DL ontology model"],
        ["POST", "/api/v1/projects/{id}/ontology/classes", "Ontology Engine", "Creates custom ontology concept class"],
        ["PUT", "/api/v1/projects/{id}/ontology/classes/{name}", "Ontology Engine", "Updates ontology class details and properties"],
        ["POST", "/api/v1/projects/{id}/ontology/export", "Ontology Engine", "Exports ontology in Turtle or OWL/XML format"],
        ["POST", "/api/v1/ontology/parse-preview", "Sandbox Viewer", "Stateless in-memory parser for raw RDF/Turtle strings"],
        ["POST", "/api/v1/ontology/upload-preview", "Sandbox Viewer", "Stateless multi-format file upload parser"],
        ["GET", "/api/v1/projects/{id}/graph/generate", "Knowledge Graph", "Generates property graph nodes and lineage edges"],
        ["POST", "/api/v1/projects/{id}/graph/export", "Knowledge Graph", "Exports graph to Cypher DDL or GraphML XML"],
        ["POST", "/api/v1/projects/{id}/graph/sync-to-target", "Knowledge Graph", "Directly executes sync to target Neo4j database"],
        ["POST", "/api/v1/projects/{id}/workflows", "Workflows", "Creates automated transformation pipeline"],
        ["GET", "/api/v1/dashboard/stats", "Dashboard", "System-wide metrics on projects, tables, rules & graphs"]
    ]
    create_styled_table(doc, api_endpoints[0], api_endpoints[1:], col_widths=[0.9, 2.3, 1.3, 2.0])

    # =========================================================================
    # SECTION 6: DATABASE SCHEMA & DATA MODEL REFERENCE
    # =========================================================================
    doc.add_page_break()
    p6 = doc.add_paragraph()
    style_heading_1(p6, "6. Database Schema & Data Model Reference")

    add_body_p(doc, "OntoForge manages application persistence and project isolation using an internal relational database (SQLite for local environments or PostgreSQL for high-availability enterprise clustering).")

    schema_tables = [
        ["Table Name", "Primary Key", "Key Foreign Keys", "Description & Stored Entities"],
        ["users", "id (UUID)", "None", "User credentials, hashed passwords, full names, active & admin flags"],
        ["roles", "id (UUID)", "None", "System access roles (ADMIN, DATA_ENGINEER, ONTOLOGIST, ANALYST)"],
        ["user_roles", "user_id, role_id", "users.id, roles.id", "Many-to-many junction mapping users to security roles"],
        ["projects", "id (UUID)", "owner_id -> users.id", "Enterprise project boundaries, unique codes, status & metadata"],
        ["source_connections", "id (UUID)", "project_id -> projects.id", "RDBMS connector definitions, host, port, db name, encrypted credentials"],
        ["graph_configs", "id (UUID)", "project_id -> projects.id", "Target knowledge graph configurations (Neo4j, Memgraph, AGE, Neptune)"],
        ["ontology_configs", "id (UUID)", "project_id -> projects.id", "Ontology naming settings, Base IRI, namespace prefixes, version strings"],
        ["metadata_tables", "id (UUID)", "project_id, source_connection_id", "Discovered relational tables, schema names, row counts, inferred domains"],
        ["metadata_columns", "id (UUID)", "table_id -> metadata_tables.id", "Physical column definitions, data types, nullability, PK/FK flags, PII tags"],
        ["profiling_results", "id (UUID)", "metadata_catalog_id", "Statistical column distributions, distinct counts, quality completeness scores"],
        ["ontology_classes", "id (UUID)", "project_id, mapped_table_id", "Semantic OWL classes, subclass taxonomy references, domain classifications"],
        ["ontology_attributes", "id (UUID)", "class_id, mapped_column_id", "OWL Datatype and Object properties, range datatypes, inverse property names"],
        ["business_rules", "id (UUID)", "project_id -> projects.id", "Governance rules, rule types, plain-English definitions, target bindings"],
        ["workflows", "id (UUID)", "project_id -> projects.id", "Automated execution pipelines, step definitions, cron scheduling expressions"],
        ["job_executions", "id (UUID)", "workflow_id -> workflows.id", "Pipeline execution history, execution timestamps, execution logs & metrics"]
    ]
    create_styled_table(doc, schema_tables[0], schema_tables[1:], col_widths=[1.5, 1.1, 1.6, 2.3])

    # =========================================================================
    # SECTION 7: DEPLOYMENT & MAINTENANCE
    # =========================================================================
    doc.add_page_break()
    p7 = doc.add_paragraph()
    style_heading_1(p7, "7. Production Deployment & Operational Maintenance")

    p7_1 = doc.add_paragraph()
    style_heading_2(p7_1, "7.1 Local Development & Testing")
    add_body_p(doc, "To launch OntoForge in a local development environment:")
    add_code_block(doc, 
        "# 1. Activate Python virtual environment\n"
        ".venv\\Scripts\\Activate.ps1  # Windows PowerShell\n"
        "# source .venv/bin/activate  # Linux/macOS\n\n"
        "# 2. Install dependencies\n"
        "pip install -r requirements.txt\n\n"
        "# 3. Launch FastAPI Server with Uvicorn\n"
        "python -m uvicorn backend.app.main:app --host 127.0.0.1 --port 8000 --reload\n\n"
        "# 4. Run automated test suite\n"
        "pytest backend/tests/\n"
    )

    p7_2 = doc.add_paragraph()
    style_heading_2(p7_2, "7.2 Production Container Deployment (Docker)")
    add_body_p(doc, "For high-availability enterprise environments, OntoForge can be containerized using Docker and orchestrated via Kubernetes or Docker Compose:")
    add_code_block(doc,
        "FROM python:3.11-slim\n"
        "WORKDIR /app\n"
        "COPY requirements.txt .\n"
        "RUN pip install --no-cache-dir -r requirements.txt\n"
        "COPY backend/ ./backend/\n"
        "EXPOSE 8000\n"
        "CMD [\"uvicorn\", \"backend.app.main:app\", \"--host\", \"0.0.0.0\", \"--port\", \"8000\", \"--workers\", \"4\"]\n"
    )

    p7_3 = doc.add_paragraph()
    style_heading_2(p7_3, "7.3 Operational Troubleshooting & FAQ")
    
    faq_data = [
        ["Symptom / Scenario", "Root Cause Analysis", "Recommended Resolution"],
        ["Connection test fails for remote PostgreSQL / SQL Server", "Firewall blockage, invalid credentials, or SSL mode mismatch", "Verify network security group rules, ensure DB allows remote connections, check user permissions."],
        ["Ontology generation produces generic relatesTo relationships", "Foreign key constraints were not explicitly defined in the source DB", "Ensure column naming adheres to standard foreign key patterns (e.g. customer_id) or manually define object properties in the editor."],
        ["Cytoscape graph rendering appears clustered / overlapping", "Physics simulation requires stabilization for large graph models", "Select the 'CoSE' layout engine from the toolbar or click 'Fit View' to automatically normalize coordinate distribution."],
        ["Stateless Sandbox fails to parse external RDF file", "File encoding issue or syntax error in external RDF/Turtle file", "Ensure file is UTF-8 encoded and conforms to standard W3C Turtle or OWL/XML serialization rules."]
    ]
    create_styled_table(doc, faq_data[0], faq_data[1:], col_widths=[1.8, 2.2, 2.5])

    # Save to file
    output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "OntoForge_Enterprise_Documentation.docx")
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

    # Also save a copy with the descriptive name
    output_path_2 = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "Quick_Pasteur_Application_Architecture_and_Usage_Guide.docx")
    doc.save(output_path_2)
    print(f"Document successfully created at: {output_path_2}")

if __name__ == "__main__":
    create_document()
